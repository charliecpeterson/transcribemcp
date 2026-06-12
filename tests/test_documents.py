"""Tests for document ingest, chunking, and retrieval tools."""
import pytest

from meetingtool import db as db_mod
from meetingtool.documents import chunk_text
from meetingtool.tools.documents import (
    add_document,
    delete_document,
    get_document,
    list_documents,
    search_documents,
)


def _mkproject(conn) -> str:
    pid = db_mod.new_id()
    conn.execute(
        "INSERT INTO projects(id, name, created_at) VALUES (?,?,?)",
        (pid, "P", db_mod.now_iso()),
    )
    return pid


def _mkmeeting(conn, pid) -> str:
    mid = db_mod.new_id()
    conn.execute(
        "INSERT INTO meetings(id, project_id, title, audio_path, status, created_at) "
        "VALUES (?,?,?,?,?,?)",
        (mid, pid, "M", "/tmp/x.wav", "ready", db_mod.now_iso()),
    )
    return mid


def test_chunk_text_respects_paragraphs():
    text = "Para one is short.\n\nPara two is also short.\n\nPara three."
    chunks = chunk_text(text, target=1000)
    assert len(chunks) == 1
    assert "Para one" in chunks[0]
    assert "Para three" in chunks[0]


def test_chunk_text_splits_near_target():
    paras = [f"Paragraph number {i} " + "x" * 200 for i in range(10)]
    text = "\n\n".join(paras)
    chunks = chunk_text(text, target=400, max_size=600)
    assert len(chunks) > 1
    assert all(len(c) <= 600 for c in chunks)


def test_chunk_text_splits_oversized_paragraph():
    giant = "A." * 1000  # 2000 chars, no paragraph breaks
    chunks = chunk_text(giant, target=400, max_size=600)
    assert len(chunks) > 1
    assert all(len(c) <= 600 for c in chunks)


def test_add_document_txt(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "notes.txt"
    f.write_text("Alice leads hiring.\n\nBob handles compensation bands.\n")
    out = add_document(pid, "Hiring Notes", str(f))
    assert out["kind"] == "txt"
    assert out["chunk_count"] >= 1
    assert out["char_count"] == len(f.read_text())


def test_add_document_md(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "plan.md"
    f.write_text("# Plan\n\nPhase one.\n")
    out = add_document(pid, "Plan", str(f))
    assert out["kind"] == "md"


def test_add_document_docx(conn, tmp_path):
    import docx

    pid = _mkproject(conn)
    f = tmp_path / "notes.docx"
    d = docx.Document()
    d.add_paragraph("Sarah Chen is the hiring manager.")
    d.add_paragraph("Diego Reyes handles the budget review.")
    d.save(str(f))
    out = add_document(pid, "Team Notes", str(f))
    assert out["kind"] == "docx"
    assert out["chunk_count"] >= 1


def test_add_document_relative_rejected(conn, tmp_path):
    pid = _mkproject(conn)
    with pytest.raises(ValueError, match="absolute"):
        add_document(pid, "x", "relative/path.txt")


def test_add_document_missing_file(conn, tmp_path):
    pid = _mkproject(conn)
    with pytest.raises(FileNotFoundError):
        add_document(pid, "x", str(tmp_path / "nope.txt"))


def test_add_document_unsupported_type(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "audio.mp3"
    f.write_bytes(b"fake")
    with pytest.raises(ValueError, match="unsupported"):
        add_document(pid, "x", str(f))


def test_add_document_unknown_project(conn, tmp_path):
    f = tmp_path / "x.txt"
    f.write_text("hi")
    with pytest.raises(ValueError, match="unknown project_id"):
        add_document("does-not-exist", "x", str(f))


def test_list_documents_filters(conn, tmp_path):
    pid = _mkproject(conn)
    pid2 = _mkproject(conn)
    mid = _mkmeeting(conn, pid)
    f = tmp_path / "a.txt"
    f.write_text("content a")
    g = tmp_path / "b.txt"
    g.write_text("content b")
    h = tmp_path / "c.txt"
    h.write_text("content c")
    add_document(pid, "A", str(f), meeting_id=mid)
    add_document(pid, "B", str(g))
    add_document(pid2, "C", str(h))

    assert len(list_documents()) == 3
    assert len(list_documents(project_id=pid)) == 2
    assert len(list_documents(meeting_id=mid)) == 1
    assert list_documents(meeting_id=mid)[0]["title"] == "A"


def test_get_document_meta(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "x.txt"
    f.write_text("hello world")
    out = add_document(pid, "X", str(f))
    meta = get_document(out["id"])
    assert meta["title"] == "X"
    assert "text" not in meta and "chunks" not in meta


def test_get_document_text_concats_chunks(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "x.txt"
    paras = ["Para " + "x" * 500 for _ in range(4)]
    f.write_text("\n\n".join(paras))
    out = add_document(pid, "X", str(f))
    doc = get_document(out["id"], format="text")
    assert "Para" in doc["text"]
    # Should not add much extra noise over the original
    assert abs(len(doc["text"]) - out["char_count"]) < 50


def test_get_document_text_truncates(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "x.txt"
    f.write_text("abcdef" * 200)
    out = add_document(pid, "X", str(f))
    doc = get_document(out["id"], format="text", max_chars=50)
    assert doc.get("truncated") is True
    assert len(doc["text"]) <= 50 + len("\n... [truncated]") + 1


def test_get_document_chunks_by_ord(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "x.txt"
    paras = [f"Paragraph {i} " + "y" * 500 for i in range(5)]
    f.write_text("\n\n".join(paras))
    out = add_document(pid, "X", str(f))
    doc = get_document(out["id"], format="chunks", chunk_ords=[0, 2])
    assert len(doc["chunks"]) == 2
    assert [c["ord"] for c in doc["chunks"]] == [0, 2]


def test_search_documents_basic(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "notes.txt"
    f.write_text(
        "Sarah Chen leads hiring.\n\n"
        "Diego Reyes reviews compensation bands quarterly.\n\n"
        "Engineering prioritizes the platform migration."
    )
    add_document(pid, "Notes", str(f))
    out = search_documents("compensation")
    assert out["count"] == 1
    assert "<<compensation>>" in out["hits"][0]["snippet"].lower() \
        or "compensation" in out["hits"][0]["snippet"].lower()


def test_search_documents_scope(conn, tmp_path):
    pid1 = _mkproject(conn)
    pid2 = _mkproject(conn)
    f = tmp_path / "a.txt"
    f.write_text("budget review")
    g = tmp_path / "b.txt"
    g.write_text("budget forecast")
    add_document(pid1, "A", str(f))
    add_document(pid2, "B", str(g))
    out = search_documents("budget", project_id=pid1)
    assert out["count"] == 1
    assert out["hits"][0]["project_id"] == pid1


def test_search_documents_empty_query_rejected(conn):
    with pytest.raises(ValueError):
        search_documents("  ")


def test_delete_document_cascades_chunks(conn, tmp_path):
    pid = _mkproject(conn)
    f = tmp_path / "x.txt"
    f.write_text("some content here that's searchable")
    out = add_document(pid, "X", str(f))
    delete_document(out["id"])
    assert search_documents("content")["count"] == 0
    assert conn.execute(
        "SELECT COUNT(*) FROM document_chunks WHERE document_id=?", (out["id"],)
    ).fetchone()[0] == 0


def test_meeting_delete_sets_document_meeting_id_null(conn, tmp_path):
    pid = _mkproject(conn)
    mid = _mkmeeting(conn, pid)
    f = tmp_path / "x.txt"
    f.write_text("agenda item")
    out = add_document(pid, "X", str(f), meeting_id=mid)
    conn.execute("DELETE FROM meetings WHERE id=?", (mid,))
    row = conn.execute(
        "SELECT meeting_id FROM documents WHERE id=?", (out["id"],)
    ).fetchone()
    assert row["meeting_id"] is None
